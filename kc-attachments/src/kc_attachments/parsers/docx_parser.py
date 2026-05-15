from __future__ import annotations
from pathlib import Path
from typing import Any

from docx import Document

from . import ParseResult, REGISTRY


def _heading_level(style_name: str) -> int | None:
    """Returns markdown heading level for Word's 'Heading N' styles, else None."""
    if not style_name:
        return None
    s = style_name.strip()
    if s.startswith("Heading "):
        try:
            return int(s.split()[1])
        except (IndexError, ValueError):
            return None
    return None


class DocxParser:
    """Walks paragraphs + tables; emits markdown headings, paragraphs, tables."""

    def parse(self, source: Path, meta: dict[str, Any]) -> ParseResult:
        doc = Document(str(source))
        out: list[str] = []
        body = doc.element.body
        from docx.oxml.ns import qn
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                for p in doc.paragraphs:
                    if p._element is child:
                        level = _heading_level(p.style.name)
                        text = p.text.strip()
                        if not text:
                            continue
                        if level:
                            out.append(f"{'#' * level} {text}")
                        else:
                            out.append(text)
                        break
            elif child.tag == qn("w:tbl"):
                for t in doc.tables:
                    if t._element is child:
                        rows = [
                            [cell.text.strip() for cell in row.cells]
                            for row in t.rows
                        ]
                        if not rows:
                            break
                        header = "| " + " | ".join(rows[0]) + " |"
                        sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                        body_rows = [
                            "| " + " | ".join(r) + " |"
                            for r in rows[1:]
                        ]
                        out.append("\n".join([header, sep, *body_rows]))
                        break
        return ParseResult(markdown="\n\n".join(out), extra_meta={})


REGISTRY["application/vnd.openxmlformats-officedocument.wordprocessingml.document"] = DocxParser()
